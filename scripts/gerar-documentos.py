from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parent.parent
OUT = ROOT / "docs" / "entregaveis"
IMG = ROOT / "qa-artifacts"
OUT.mkdir(parents=True, exist_ok=True)

VERDE = "075B35"
VERDE_CLARO = "E8F5EE"
AMARELO = "D9ED29"
AZUL = "2F67E8"
CINZA = "526680"
TEXTO = "12253F"

def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")

def add_field(paragraph, instruction):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "Atualize os campos no Word"
    def field_run(element):
        run = OxmlElement("w:r")
        run.append(element)
        return run
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    paragraph._p.append(field_run(begin))
    paragraph._p.append(field_run(instr))
    paragraph._p.append(field_run(separate))
    paragraph._p.append(field_run(text))
    paragraph._p.append(field_run(end))

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.add_run("Página ")
    add_field(paragraph, "PAGE")

def base_document():
    d = Document()
    sec = d.sections[0]
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)
    styles = d.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor.from_string(TEXTO)
    for style_name, size, color in (("Title", 28, VERDE), ("Heading 1", 18, VERDE), ("Heading 2", 13, AZUL), ("Heading 3", 11, VERDE)):
        s = styles[style_name]
        s.font.name = "Aptos Display" if style_name == "Title" else "Aptos"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)
    footer = sec.footer.paragraphs[0]
    footer.text = "Comitê João Jorge  |  Documento controlado"
    footer.style = styles["Normal"]
    footer.runs[0].font.size = Pt(8)
    add_page_number(sec.footer.add_paragraph())
    settings = d.settings.element
    update = OxmlElement("w:updateFields")
    update.set(qn("w:val"), "true")
    settings.append(update)
    return d

def cover(d, title, subtitle):
    for _ in range(5):
        d.add_paragraph()
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("COMITÊ JOÃO JORGE")
    r.bold = True; r.font.size = Pt(15); r.font.color.rgb = RGBColor.from_string(VERDE)
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title)
    r.bold = True; r.font.size = Pt(29); r.font.color.rgb = RGBColor.from_string(VERDE)
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(subtitle)
    r.font.size = Pt(14); r.font.color.rgb = RGBColor.from_string(CINZA)
    d.add_paragraph()
    box = d.add_table(rows=4, cols=2)
    box.alignment = WD_TABLE_ALIGNMENT.CENTER
    box.autofit = False
    labels = [("Responsável", "Rafael Santos de Menezes"), ("Versão", "1.0"), ("Emissão", "13 de agosto de 2026"), ("Classificação", "Documento interno")]
    for row, (k, v) in zip(box.rows, labels):
        row.cells[0].width = Cm(4.2); row.cells[1].width = Cm(9.0)
        set_cell_shading(row.cells[0], VERDE); set_cell_shading(row.cells[1], VERDE_CLARO)
        for c in row.cells: set_cell_margins(c); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        rk = row.cells[0].paragraphs[0].add_run(k); rk.bold = True; rk.font.color.rgb = RGBColor(255,255,255)
        rv = row.cells[1].paragraphs[0].add_run(v); rv.font.color.rgb = RGBColor.from_string(TEXTO)
    d.add_page_break()

def version_table(d):
    d.add_heading("Controle de versão", level=1)
    table = d.add_table(rows=2, cols=5)
    table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Versão", "Data", "Responsável", "Descrição", "Status"]
    values = ["1.0", "13 de agosto de 2026", "Rafael Santos de Menezes", "Emissão inicial", "Final"]
    for i, txt in enumerate(headers):
        set_cell_shading(table.rows[0].cells[i], VERDE)
        run = table.rows[0].cells[i].paragraphs[0].add_run(txt); run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
    for i, txt in enumerate(values):
        set_cell_shading(table.rows[1].cells[i], "F4F7F9")
        table.rows[1].cells[i].paragraphs[0].add_run(txt)
    d.add_paragraph()

def index_pages(d):
    d.add_heading("Sumário", level=1)
    p = d.add_paragraph()
    add_field(p, 'TOC \\o "1-3" \\h \\z \\u')
    d.add_page_break()
    d.add_heading("Índice de figuras", level=1)
    p = d.add_paragraph()
    add_field(p, 'TOC \\h \\z \\c "Figura"')
    d.add_paragraph("Observação: ao abrir no Microsoft Word, use Ctrl+A e F9 caso seja necessário atualizar o sumário e o índice de figuras.")
    d.add_page_break()

def legend(d):
    d.add_heading("Legenda", level=1)
    items = [
        ("Tela desktop", "captura em resolução 1440 × 900."),
        ("Tela celular", "captura em resolução 390 × 844."),
        ("Imagem de QA", "evidência registrada pelo roteiro automatizado local."),
        ("Valores ocultos", "informações financeiras protegidas pelo estado de privacidade do Painel."),
    ]
    table = d.add_table(rows=0, cols=2); table.style = "Table Grid"
    for a, b in items:
        cells = table.add_row().cells
        set_cell_shading(cells[0], VERDE_CLARO)
        cells[0].paragraphs[0].add_run(a).bold = True
        cells[1].paragraphs[0].add_run(b)
    d.add_page_break()

def caption(d, text):
    p = d.add_paragraph(style="Caption")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run("Figura ")
    add_field(p, "SEQ Figura \\* ARABIC")
    p.add_run(". " + text)

def add_image(d, filename, text, width=15.6):
    path = IMG / filename
    if not path.exists():
        return
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(path), width=Cm(width))
    caption(d, text)

def add_pair(d, left_file, right_file, text):
    table = d.add_table(rows=1, cols=2); table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for cell, filename in zip(table.rows[0].cells, (left_file, right_file)):
        cell.width = Cm(7.6); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell, 60, 60, 60, 60)
        p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(IMG / filename), width=Cm(6.9))
    caption(d, text)

def para(d, text):
    p = d.add_paragraph(text)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.line_spacing = 1.15
    return p

def bullets(d, values):
    for value in values:
        d.add_paragraph(value, style="List Bullet")

def flow_report():
    d = base_document()
    cover(d, "Relatório de Fluxos do Sistema", "Comitê João Jorge")
    version_table(d); index_pages(d); legend(d)
    d.add_heading("1. Apresentação", level=1)
    para(d, "Este relatório consolida os fluxos funcionais do sistema Comitê João Jorge. O conteúdo foi elaborado a partir do manual operacional, da análise da aplicação, das evidências de QA local e das telas registradas em desktop e celular.")
    d.add_heading("2. Objetivo e escopo", level=1)
    para(d, "O sistema centraliza o controle da equipe de campanha, da reserva de candidatos a contratação, da presença, dos valores entregues, das entradas de caixa e da folha de pagamento. A aplicação é responsiva e foi projetada para uso contínuo em computador e telefone.")
    d.add_heading("3. Fluxo de acesso e sessão", level=1)
    bullets(d, ["O operador informa e-mail e senha na tela de acesso.", "Após validação, o Painel é exibido como tela inicial.", "A opção Sair do sistema encerra a sessão e restaura a privacidade dos valores financeiros.", "O sistema preserva a escolha de exibir ou ocultar os valores enquanto a sessão estiver ativa."])
    d.add_heading("4. Fluxo do Painel", level=1)
    para(d, "O Painel apresenta Entradas, Saídas, pessoas Ativas e pessoas Sem assinatura. Os cards de entradas e saídas permitem consultar o histórico financeiro e exportar planilhas. Os cards Ativos e Sem assinatura direcionam o operador para a gestão da Equipe.")
    add_pair(d, "desktop-loop-1-painel.png", "mobile-loop-1-painel.png", "Painel nas versões desktop, à esquerda, e celular, à direita.")
    d.add_heading("5. Fluxo da Equipe e Reserva", level=1)
    para(d, "Cada pessoa possui cadastro individual com identificação, perfil, equipe, data de entrada, telefone opcional e situação da assinatura. Pessoas ativas permanecem no topo da listagem, reservas ficam na faixa intermediária e desligados aparecem ao final. O histórico é preservado porque não existe exclusão física da pessoa.")
    add_image(d, "desktop-loop-1-equipe.png", "Tela de Equipe com filtros e acesso aos cadastros.")
    add_image(d, "desktop-loop-1-novo-funcionario.png", "Cadastro de funcionário com dados de identificação, perfil, equipe e assinatura.")
    d.add_heading("6. Fluxo de Valores", level=1)
    para(d, "O operador abre Registrar Pagamento, busca a pessoa pelo nome, seleciona o tipo de valor e informa o montante. O histórico apresenta a data do lançamento e a tabela exportada inclui os valores em BRL e o total geral.")
    add_pair(d, "desktop-loop-1-modal-registrar-pagamento.png", "mobile-loop-1-valores.png", "Registro de pagamento no desktop e consulta de Valores no celular.")
    d.add_heading("7. Fluxo de Salário", level=1)
    para(d, "A nova lista de pagamento inicia com a data vigente. Cada funcionário apresenta o valor previsto, o estado de assinatura e o controle Pago. O card da pessoa abre sua ficha para edição. Antes de registrar, o operador pode consultar o total da folha, o total já pago e, opcionalmente, o saldo disponível e restante.")
    para(d, "Ao tentar sair da tela antes do registro, o sistema solicita confirmação. Isso evita perda de valores preenchidos por engano.")
    add_image(d, "desktop-loop-1-nova-lista-salario.png", "Nova lista de pagamento na versão desktop.")
    add_image(d, "desktop-loop-1-aviso-saida-salario.png", "Confirmação de saída da lista de pagamento sem registro.")
    d.add_heading("8. Fluxo de Presença", level=1)
    para(d, "A presença é lançada por data. A ação Marcar todos estabelece o cenário inicial e o operador remove somente os ausentes. A aplicação impede o lançamento duplicado da mesma data. A saída sem lançar também exige confirmação.")
    add_pair(d, "desktop-loop-1-presenca-marcada.png", "mobile-loop-1-nova-lista-presenca.png", "Lançamento de presença no desktop e na versão celular.")
    d.add_heading("9. Integridade, sincronização e exportação", level=1)
    bullets(d, ["Presença e salário utilizam operações atômicas no banco, evitando listas parcialmente gravadas.", "Atualizações de pessoas, valores, caixa, presença e pagamentos são acompanhadas em tempo real.", "Em edição simultânea do mesmo registro, o último salvamento prevalece. A orientação operacional é evitar edição concorrente da mesma informação.", "Todas as planilhas incluem valores em BRL e totalização final."])
    d.add_heading("10. Responsividade e validação", level=1)
    para(d, "O roteiro local executou três ciclos em 1440 × 900 e três ciclos em 390 × 844. Foram verificados login, Painel, Valores, Salário, confirmação de saída, Presença, marcação em massa, Equipe e novo funcionário. Os ciclos foram concluídos sem alteração dos dados existentes.")
    d.add_heading("11. Conclusão", level=1)
    para(d, "O sistema apresenta fluxos consistentes para a rotina de campanha, com foco em proteção do histórico, agilidade em operações repetitivas e utilização em dispositivos móveis. As confirmações de saída, a segmentação entre reserva, ativos e desligados e as exportações fortalecem a rastreabilidade operacional.")
    d.save(OUT / "Relatorio_de_Fluxos_Comite_Joao_Jorge.docx")

def photo_report():
    d = base_document()
    cover(d, "Relatório Fotográfico", "Comitê João Jorge")
    version_table(d); index_pages(d); legend(d)
    d.add_heading("1. Finalidade", level=1)
    para(d, "Este relatório reúne as evidências visuais do roteiro de QA local. As capturas registram a mesma aplicação em desktop e celular, com foco na disposição dos elementos, navegação e estados críticos.")
    d.add_heading("2. Painel", level=1)
    add_pair(d, "desktop-loop-1-painel.png", "mobile-loop-1-painel.png", "Painel com cards financeiros, indicadores de equipe e gráfico de gastos.")
    d.add_heading("3. Valores", level=1)
    add_pair(d, "desktop-loop-1-valores.png", "mobile-loop-1-valores.png", "Consulta de Valores lançados em desktop e celular.")
    add_image(d, "desktop-loop-1-modal-registrar-pagamento.png", "Modal Registrar Pagamento com pesquisa de pessoa e tipo de lançamento.")
    d.add_heading("4. Salário", level=1)
    add_pair(d, "desktop-loop-1-nova-lista-salario.png", "mobile-loop-1-nova-lista-salario.png", "Nova lista de pagamento nas duas responsividades.")
    add_image(d, "desktop-loop-1-aviso-saida-salario.png", "Proteção contra saída da tela de pagamento com dados não registrados.")
    d.add_heading("5. Presença", level=1)
    add_pair(d, "desktop-loop-1-presenca-marcada.png", "mobile-loop-1-nova-lista-presenca.png", "Lançamento de presença com marcação de participantes.")
    d.add_heading("6. Equipe", level=1)
    add_image(d, "desktop-loop-1-equipe.png", "Tela de Equipe, filtros e organização da lista de pessoas.")
    add_image(d, "desktop-loop-1-novo-funcionario.png", "Tela de cadastro de funcionário.")
    d.add_heading("7. Registro de execução", level=1)
    para(d, "Foram concluídos três ciclos de navegação em desktop e três em celular. As imagens foram produzidas no ambiente local sem cadastro, edição ou remoção de dados do sistema.")
    d.save(OUT / "Relatorio_Fotografico_Comite_Joao_Jorge.docx")

def manual():
    d = base_document()
    cover(d, "Manual Completo do Sistema", "Comitê João Jorge")
    version_table(d); index_pages(d); legend(d)
    d.add_heading("1. Como entrar e sair", level=1)
    para(d, "Abra o sistema, informe seu e-mail e senha e selecione Entrar. Para encerrar, use Sair do sistema no cabeçalho. O sistema restaura a ocultação dos valores financeiros quando a sessão termina.")
    d.add_heading("2. Painel", level=1)
    para(d, "O Painel é a tela inicial. Entradas mostra os recursos recebidos em caixa. Saídas reúne valores entregues e salários pagos. Toque em qualquer um desses cards para consultar o histórico e exportar a planilha com total. Ativos e Sem assinatura abrem a tela de Equipe.")
    para(d, "O ícone de olho mostra ou esconde valores. A escolha permanece entre as telas enquanto o usuário não sair do sistema.")
    add_pair(d, "desktop-loop-1-painel.png", "mobile-loop-1-painel.png", "Painel atualizado nas versões desktop e celular.")
    d.add_heading("3. Equipe, reserva e cadastro", level=1)
    para(d, "Use Novo para cadastrar uma pessoa. Informe nome, CPF ou RG, salário, endereço, perfil, equipe, situação de assinatura, data de entrada e, se necessário, telefone. O telefone é opcional e aceita formato de telefone fixo ou celular.")
    para(d, "A situação pode ser Ativa, Reserva ou Desligada. Reservas ficam disponíveis para ativação futura. Desligados não entram nas novas rotinas, mas preservam o histórico. Toque no card da pessoa para abrir a edição. Use os filtros de nome, perfil, equipe e assinatura para localizar registros.")
    add_image(d, "desktop-loop-1-equipe.png", "Tela de Equipe com filtros de consulta.")
    add_image(d, "desktop-loop-1-novo-funcionario.png", "Cadastro atualizado de funcionário.")
    d.add_heading("4. Valores", level=1)
    para(d, "Na tela Valores, selecione Novo. O modal se chama Registrar Pagamento. Pesquise a pessoa pelo nome, informe o valor e escolha o tipo de lançamento. Vale da equipe, Combustível, Alimentação e Outro são as opções disponíveis.")
    para(d, "O histórico apresenta a data de cada lançamento. A opção Tabela permite exportar a planilha, que inclui os lançamentos em BRL e o total geral.")
    add_pair(d, "desktop-loop-1-modal-registrar-pagamento.png", "mobile-loop-1-valores.png", "Registro de pagamento e histórico de Valores.")
    d.add_heading("5. Salário", level=1)
    para(d, "Use Nova para criar uma lista de pagamento. O período começa com a data atual. Informe o valor de cada pessoa e toque em Pagar para marcar o recebimento. O card da pessoa abre a ficha de edição, enquanto o botão Pagar controla apenas o estado de pagamento.")
    para(d, "O sistema mostra Total da folha, Já pago e Saldo total opcional. Após conferir, selecione Registrar. Se sair antes, será exibido um aviso. A tabela de histórico mostra também a situação da assinatura.")
    add_pair(d, "desktop-loop-1-nova-lista-salario.png", "mobile-loop-1-nova-lista-salario.png", "Nova lista de pagamento em desktop e celular.")
    add_image(d, "desktop-loop-1-aviso-saida-salario.png", "Aviso de proteção contra perda de dados na tela de salário.")
    d.add_heading("6. Presença", level=1)
    para(d, "Selecione Nova para iniciar uma lista. A data atual vem preenchida. Use Marcar todos e desmarque apenas quem faltou. Em seguida, selecione Lançar e confirme. A mesma data não pode ser registrada duas vezes.")
    para(d, "A Tabela permite consultar e exportar o período. O menu de três pontos em cada lista permite excluir o lançamento, quando necessário. Ao sair antes de lançar, confirme se deseja perder as marcações.")
    add_pair(d, "desktop-loop-1-presenca-marcada.png", "mobile-loop-1-nova-lista-presenca.png", "Lançamento de presença na versão atual.")
    d.add_heading("7. Planilhas, conexão e operação simultânea", level=1)
    bullets(d, ["As planilhas de Entradas, Saídas, Valores, Salário e Presença incluem totalização e valores em BRL.", "Se a conexão cair, aguarde o retorno antes de repetir qualquer salvamento.", "Duas pessoas podem acompanhar o sistema simultaneamente. Evite editar o mesmo registro ao mesmo tempo, pois o último envio prevalece.", "Presença e pagamento são gravados por operações atômicas, reduzindo o risco de listas incompletas."])
    d.add_heading("8. Cuidados operacionais", level=1)
    bullets(d, ["Não compartilhe a senha fora da equipe autorizada.", "Confirme os valores antes de Registrar um pagamento.", "Mantenha o cadastro de assinatura atualizado antes de processar a folha.", "Use Sair do sistema ao finalizar o trabalho em aparelho compartilhado."])
    d.save(OUT / "Manual_Completo_Comite_Joao_Jorge.docx")

def proposal():
    d = base_document()
    cover(d, "Proposta de Solução", "Governança para Comitê de Campanha")
    version_table(d); index_pages(d); legend(d)
    d.add_heading("1. Resumo executivo", level=1)
    para(d, "Governança é uma solução digital para organizar a rotina operacional de um comitê político durante a campanha. Em uma única aplicação, a direção acompanha equipe, reservas, assinaturas, presença, entradas de caixa, valores entregues e folha de pagamento. O projeto prioriza agilidade em campo, rastreabilidade financeira e uso simples em celular.")
    d.add_heading("2. A solução", level=1)
    para(d, "A aplicação foi desenhada para servir como centro de controle da campanha. A equipe administrativa registra informações no momento em que elas acontecem, enquanto a direção consulta os indicadores consolidados no Painel. A navegação se adapta a computador, notebook e telefone, sem exigir instalação no aparelho.")
    add_pair(d, "desktop-loop-1-painel.png", "mobile-loop-1-painel.png", "Identidade visual da solução Governança em desktop e celular.")
    d.add_heading("3. Inclusões da solução", level=1)
    table = d.add_table(rows=1, cols=2); table.style = "Table Grid"; table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, text in enumerate(("Entrega", "Descrição")):
        set_cell_shading(table.rows[0].cells[i], VERDE)
        run = table.rows[0].cells[i].paragraphs[0].add_run(text); run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
    inclusoes = [
        ("Painel executivo", "Indicadores de entradas, saídas, equipe ativa, documentos pendentes e distribuição de gastos."),
        ("Gestão de equipe", "Cadastro, filtros, status ativo, reserva e desligado, com preservação do histórico."),
        ("Controle de presença", "Listas por data, marcação em massa, validação contra duplicidade e exportação."),
        ("Gestão financeira", "Entradas de caixa, valores entregues, folha de pagamento, saldos e planilhas totalizadas."),
        ("Documentação", "Manual completo, relatório de fluxos, relatório fotográfico e relatório de QA local."),
        ("Responsividade", "Experiência adaptada para desktop, notebook e celular."),
    ]
    for a, b in inclusoes:
        cells = table.add_row().cells; set_cell_shading(cells[0], VERDE_CLARO)
        cells[0].paragraphs[0].add_run(a).bold = True; cells[1].paragraphs[0].add_run(b)
    d.add_heading("4. Funcionalidades detalhadas", level=1)
    d.add_heading("4.1 Gestão de pessoas", level=2)
    bullets(d, ["Cadastro de colaboradores com dados de identificação, telefone opcional, perfil, equipe e data de entrada.", "Organização entre ativos, reservas e desligados, sem perda de histórico.", "Controle de assinatura de documento e filtros para localizar rapidamente cada situação.", "Acesso à edição ao tocar no card da pessoa."])
    add_image(d, "desktop-loop-1-equipe.png", "Gestão de equipe, filtros e organização de registros.")
    d.add_heading("4.2 Valores e pagamentos", level=2)
    bullets(d, ["Registro de Vale da equipe, Combustível, Alimentação e outros pagamentos.", "Histórico por pessoa e por data de lançamento.", "Nova lista de salários com total da folha, total já pago e saldo opcional.", "Proteção contra saída involuntária de telas com preenchimento em andamento.", "Exportações com valores em BRL e total geral."])
    add_pair(d, "desktop-loop-1-modal-registrar-pagamento.png", "desktop-loop-1-nova-lista-salario.png", "Registro de pagamento e operação de folha salarial.")
    d.add_heading("4.3 Presença e acompanhamento", level=2)
    bullets(d, ["Listas de presença por dia, com data atual preenchida automaticamente.", "Ação Marcar todos para acelerar a chamada diária.", "Histórico consultável e exportável.", "Validação para impedir lançamento repetido na mesma data."])
    add_pair(d, "desktop-loop-1-presenca-marcada.png", "mobile-loop-1-nova-lista-presenca.png", "Presença em desktop e celular.")
    d.add_heading("5. Hospedagem, base de dados e continuidade", level=1)
    para(d, "A solução utiliza hospedagem Vercel e base de dados Supabase. A arquitetura atual separa a aplicação web da camada de dados, permitindo evolução controlada. A operação utiliza atualização em tempo real para que operadores conectados recebam alterações relevantes sem precisar atualizar a página.")
    bullets(d, ["Hospedagem da aplicação web em ambiente Vercel.", "Base de dados centralizada em Supabase.", "Operações atômicas para listas de presença e pagamentos, reduzindo risco de gravação incompleta.", "Exportação de planilhas para apoio à prestação de contas.", "Rotina de backup recomendada: exportação periódica do banco e retenção externa controlada pela direção."])
    para(d, "A configuração de frequência, retenção, local de armazenamento e responsável pelo backup deve ser formalizada na contratação. Para operação de campanha, recomenda-se backup diário e uma cópia externa semanal.")
    d.add_heading("6. Suporte técnico e canal de emergência", level=1)
    para(d, "A entrega pode incluir suporte técnico para orientação de uso, diagnóstico de incidentes, acompanhamento de atualizações e correção de falhas identificadas. O atendimento deve ser organizado em dois níveis para diferenciar dúvidas operacionais de indisponibilidade crítica.")
    table = d.add_table(rows=1, cols=3); table.style = "Table Grid"
    for i, text in enumerate(("Nível", "Exemplos", "Canal sugerido")):
        set_cell_shading(table.rows[0].cells[i], VERDE)
        run = table.rows[0].cells[i].paragraphs[0].add_run(text); run.bold = True; run.font.color.rgb = RGBColor(255,255,255)
    niveis = [
        ("Suporte operacional", "Dúvidas de uso, cadastro, relatórios e orientação de rotina.", "Canal de suporte definido pela direção."),
        ("Emergência", "Sistema indisponível, falha de acesso generalizada ou risco de perda de operação em dia de campanha.", "Canal prioritário a ser definido em contrato, por exemplo WhatsApp corporativo."),
    ]
    for row in niveis:
        cells = table.add_row().cells
        for i, text in enumerate(row): cells[i].paragraphs[0].add_run(text)
    para(d, "Horário de atendimento, prazo de primeira resposta, escopo de correções e responsáveis devem constar no acordo comercial. Isso assegura um canal de emergência realista e mensurável para a campanha.")
    d.add_heading("7. Treinamento e transferência de operação", level=1)
    bullets(d, ["Treinamento inicial da direção e dos operadores autorizados.", "Demonstração prática de Painel, Equipe, Valores, Salário, Presença e exportações.", "Orientação sobre reserva, desligamento, assinatura e boas práticas de registro.", "Entrega do Manual Completo do Sistema e dos relatórios técnicos.", "Recomendação de uma simulação assistida antes do início da operação de campanha."])
    d.add_heading("8. Qualidade, auditoria e segurança operacional", level=1)
    para(d, "A solução foi submetida a roteiro local de QA com ciclos em desktop e celular. Foram verificados login, navegação, valores, nova folha, confirmação de saída, presença, marcação em massa, equipe e novo cadastro. O roteiro não criou nem alterou dados da campanha durante a execução.")
    bullets(d, ["Valores financeiros ocultos por padrão no Painel.", "Histórico preservado para pessoas desligadas.", "Confirmações antes de descartar lançamentos em andamento.", "Planilhas com totalização para conferência.", "Auditoria operacional favorecida por históricos, filtros e exportações."])
    d.add_heading("9. Evoluções recomendadas", level=1)
    bullets(d, ["Política de backup automatizado e teste periódico de restauração.", "Perfis de acesso separados, evitando uso compartilhado de uma única senha.", "Trilha de auditoria com autor, data e hora das alterações sensíveis.", "Ambiente de homologação separado do ambiente de produção.", "Monitoramento de disponibilidade e alertas técnicos.", "Integração de documentos digitalizados, caso a direção deseje centralizar comprovações."])
    d.add_heading("10. Entregáveis", level=1)
    bullets(d, ["Aplicação Governança pronta para uso no endereço de hospedagem definido.", "Acesso administrativo e configuração da base de dados conforme escopo contratado.", "Manual completo atualizado.", "Relatório de fluxos do sistema.", "Relatório fotográfico de responsividade e navegação.", "Relatório de QA local.", "Treinamento inicial e plano de suporte a formalizar."])
    d.add_heading("11. Considerações comerciais", level=1)
    para(d, "Esta proposta descreve o produto e o ciclo de entrega. Valores, vigência, nível de serviço, número de usuários, canal de emergência, responsabilidade sobre backups e eventuais evoluções devem ser definidos em instrumento comercial próprio. Dessa forma, a direção recebe uma solução operacionalmente completa e um plano transparente de continuidade.")
    d.save(OUT / "Proposta_Comercial_Governanca_Comite_Joao_Jorge.docx")

flow_report()
photo_report()
manual()
proposal()
print(f"Documentos gerados em {OUT}")
